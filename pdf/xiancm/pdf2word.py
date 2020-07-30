#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jan 30 15:55:45 2020

@author: Gabriel
"""

# =============================================================================
# PDF2Word
# =============================================================================

import pdfplumber
import os

from docx import Document

# 1.遍历文件夹中的所有PDF文件
file_dir = r'/Users/Gabriel/cloud/workspace/code/financedatahub/exc/pdf/xiancm'  # 所要遍历的文件夹
file_list = []
for files in os.walk(file_dir):  # 遍历该文件夹及其里面的所有子文件夹
    for file in files[2]:
        if os.path.splitext(file)[1] == '.pdf' or os.path.splitext(file)[1] == '.PDF':
            file_list.append(file_dir + '/' + file)
            
            # file_list.append(file_dir + '\\' + file)   # Win下文件路径
print(file_list)

# 2.PDF文本解析和内容筛选
pdf_all = []
document = Document()
for i in range(len(file_list)):
    pdf = pdfplumber.open(file_list[i])
    pages = pdf.pages
    text_all = []
    for page in pages:  # 遍历pages中每一页的信息
        text = page.extract_text()  # 提取当页的文本内容
        text_all.append(text)  # 通过列表.append()方法汇总每一页内容
    text_all = ''.join(text_all)  # 把列表转换成字符串
    # print(text_all)  # 打印全部文本内容
    
    #内容存入word
    
    document.add_paragraph(text_all)


    
    pdf.close()
  
document.save('filename.docx')
print('PDF文本输出完毕')


    # 通过正文进行筛选
#     if ('自有' in text_all) or ('议案' in text_all) or ('理财' in text_all) or ('现金管理' in text_all):
#         pdf_all.append(file_list[i])
# print(pdf_all)  # 打印筛选后的PDF列表

# # 3.筛选后文件的移动
# for pdf_i in pdf_all:
#     newpath = '/Users/Gabriel/cloud/workspace/code/financedatahub/exc/pdf/筛选后的文件夹/' + pdf_i.split('/')[-1]  # 这边这个移动到的文件夹一定要提前就创建好！
#     os.rename(pdf_i, newpath)  # 执行移动操作

# print('PDF文本解析及筛选完毕！')
    
    
    
# 输出word文档
# from docx import Document
# document = Document()

# document.add_paragraph(text_all)

# document.save('filename.docx')


    
